"""DOCX 写入器 — 将 HTML 内容转换为 Word 文档"""

import logging
import re
from pathlib import Path
from typing import Optional, Union

logger = logging.getLogger(__name__)


class DOCXWriter:
    """将 HTML 内容转换为 Word (.docx) 文件。

    支持两种转换方式：
    1. 优先使用 mammoth（更高质量的 HTML → DOCX 转换）
    2. 回退到 python-docx（手动解析 HTML 标签）

    用法::

        writer = DOCXWriter()
        writer.write("<html><body><h1>Hello</h1></body></html>", "output.docx")
    """

    def __init__(self):
        self._mammoth_available = None
        self._docx_available = None

    def _check_mammoth(self) -> bool:
        if self._mammoth_available is None:
            try:
                import mammoth  # noqa: F401
                self._mammoth_available = True
            except ImportError:
                self._mammoth_available = False
        return self._mammoth_available

    def _check_docx(self) -> bool:
        if self._docx_available is None:
            try:
                import docx  # noqa: F401
                self._docx_available = True
            except ImportError:
                self._docx_available = False
        return self._docx_available

    def write(
        self,
        html_content: str,
        output_path: Union[str, Path],
    ) -> Path:
        """将 HTML 转换为 Word 文件。

        Args:
            html_content: 完整的 HTML 文档字符串
            output_path: 输出 docx 文件路径

        Returns:
            生成的 docx 文件路径

        Raises:
            ImportError: 无可用的 docx 库时
            Exception: 转换失败时
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # 优先使用 mammoth
        if self._check_mammoth():
            return self._write_with_mammoth(html_content, output_path)

        # 回退到 python-docx
        if self._check_docx():
            return self._write_with_docx(html_content, output_path)

        raise ImportError(
            "无可用的 docx 库。请安装其中之一：\n"
            "  pip install mammoth  （推荐，HTML→DOCX 转换质量高）\n"
            "  pip install python-docx  （回退方案）"
        )

    def _write_with_mammoth(self, html_content: str, output_path: Path) -> Path:
        """使用 mammoth 进行高质量 HTML→DOCX 转换"""
        try:
            import mammoth

            result = mammoth.convert_to_docx(
                html_content,
                convert_image=mammoth.images.data_uri,  # 将图片嵌入为 data URI
            )
            output_path.write_bytes(result.value)

            if result.messages:
                for msg in result.messages:
                    logger.warning("mammoth 警告: %s", msg)

            logger.info("DOCX 已生成 (mammoth): %s", output_path)
            return output_path
        except Exception as e:
            logger.error("mammoth 转换失败: %s，尝试 python-docx", e)
            if self._check_docx():
                return self._write_with_docx(html_content, output_path)
            raise

    def _write_with_docx(self, html_content: str, output_path: Path) -> Path:
        """使用 python-docx 手动解析 HTML 并生成 DOCX"""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # 设置默认字体
            style = doc.styles["Normal"]
            font = style.font
            font.size = Pt(11)
            font.name = "Microsoft YaHei"

            # 提取标题
            titles = re.findall(r"<h([1-6])[^>]*>(.*?)</h\1>", html_content, re.DOTALL)
            paragraphs = re.findall(r"<p[^>]*>(.*?)</p>", html_content, re.DOTALL)
            list_items = re.findall(r"<li[^>]*>(.*?)</li>", html_content, re.DOTALL)
            tables = re.findall(r"<table[^>]*>(.*?)</table>", html_content, re.DOTALL)
            warns = re.findall(r'<div\s+class="warning"[^>]*>(.*?)</div>', html_content, re.DOTALL)

            # 辅助：清理 HTML 标签
            clean = lambda t: re.sub(r"<[^>]+>", "", t).strip()

            # 添加主标题
            if titles:
                level, title_text = titles[0]
                doc.add_heading(clean(title_text), level=0)
            elif content_title := re.search(r"<title[^>]*>(.*?)</title>", html_content, re.DOTALL):
                doc.add_heading(clean(content_title.group(1)), level=0)

            # 添加所有段落
            for p_text in paragraphs:
                text = clean(p_text)
                if text:
                    doc.add_paragraph(text)

            # 添加列表项
            if list_items:
                for item in list_items:
                    text = clean(item)
                    if text:
                        doc.add_paragraph(text, style="List Bullet")

            # 添加表格
            for table_html in tables:
                rows = re.findall(r"<tr[^>]*>(.*?)</tr>", table_html, re.DOTALL)
                if rows:
                    cells_in_first_row = re.findall(r"<t[dh][^>]*>(.*?)</t[dh]>", rows[0], re.DOTALL)
                    num_cols = max(len(cells_in_first_row), 1)
                    table = doc.add_table(rows=len(rows), cols=num_cols)
                    table.style = "Table Grid"
                    for i, row_html in enumerate(rows):
                        cells = re.findall(r"<t[dh][^>]*>(.*?)</t[dh]>", row_html, re.DOTALL)
                        for j, cell_html in enumerate(cells):
                            if j < num_cols:
                                table.cell(i, j).text = clean(cell_html)

            # 添加警告块
            for warn in warns:
                doc.add_paragraph("")
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run("⚠ " + clean(warn))
                run.bold = True
                from docx.shared import RGBColor
                run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)

            # 如果以上都空，尝试提取所有文本
            if not titles and not paragraphs and not list_items:
                all_text = re.sub(r"<[^>]+>", "\n", html_content)
                for line in all_text.split("\n"):
                    line = line.strip()
                    if line and len(line) > 2:
                        doc.add_paragraph(line)

            doc.save(str(output_path))
            logger.info("DOCX 已生成 (python-docx): %s", output_path)
            return output_path

        except Exception as e:
            logger.error("python-docx 转换失败: %s", e)
            raise
