"""Agent Tool 定义 — 供 LangGraph Agent 调用的工具集"""

import os
from pathlib import Path
from typing import Optional

from langchain_core.tools import tool

# 项目根目录
BASE_DIR = Path(__file__).parent.parent.parent
UPLOAD_DIR = BASE_DIR / "uploads"
OUTPUT_DIR = BASE_DIR / "output"


@tool
def extract_text_from_file(file_path: str) -> str:
    """从上传的文件中提取文本内容。支持 PDF、Word(.docx)、TXT 格式。

    Args:
        file_path: 文件的绝对路径

    Returns:
        提取的文本内容
    """
    path = Path(file_path)
    if not path.exists():
        return f"错误: 文件不存在 {file_path}"

    suffix = path.suffix.lower()

    try:
        if suffix == ".pdf":
            return _extract_pdf(path)
        elif suffix in (".docx", ".doc"):
            return _extract_docx(path)
        elif suffix in (".txt", ".md", ".csv", ".json"):
            return path.read_text(encoding="utf-8", errors="replace")
        elif suffix in (".jpg", ".jpeg", ".png", ".gif", ".webp"):
            return f"[图片文件: {path.name}] 请使用 analyze_image 工具分析此图片。"
        else:
            return f"不支持的文件格式: {suffix}"
    except Exception as e:
        return f"提取文本失败: {str(e)}"


@tool
def analyze_image(image_path: str, description: str = "") -> str:
    """使用多模态 LLM 分析图片内容，返回详细的图片描述。

    Args:
        image_path: 图片文件的绝对路径
        description: 用户对图片的额外描述（可选）

    Returns:
        图片内容的详细描述
    """
    from src.agent.llm import get_llm_safe
    from langchain_core.messages import HumanMessage

    llm = get_llm_safe(temperature=0.2)
    if not llm:
        return f"[图片: {image_path}] LLM 不可用，无法分析图片内容。"

    path = Path(image_path)
    if not path.exists():
        return f"错误: 图片文件不存在 {image_path}"

    try:
        import base64
        with open(path, "rb") as f:
            image_data = base64.b64encode(f.read()).decode("utf-8")

        suffix = path.suffix.lower()
        mime_map = {".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".png": "image/png",
                    ".gif": "image/gif", ".webp": "image/webp"}
        mime_type = mime_map.get(suffix, "image/jpeg")

        prompt_text = "请详细描述这张图片的内容，包括所有可见的文字、图表、产品细节等。"
        if description:
            prompt_text += f"\n用户备注: {description}"

        message = HumanMessage(content=[
            {"type": "text", "text": prompt_text},
            {"type": "image_url", "image_url": {
                "url": f"data:{mime_type};base64,{image_data}",
            }},
        ])

        response = llm.invoke([message])
        return response.content
    except Exception as e:
        return f"图片分析失败: {str(e)}"


@tool
def generate_html(content: str, layout: str = "professional", style: str = "clean") -> str:
    """根据内容和排版要求生成 HTML 文档。

    Args:
        content: 文档内容（纯文本或 Markdown）
        layout: 布局风格（professional/casual/technical）
        style: 样式（clean/modern/minimal）

    Returns:
        完整的 HTML 文档字符串
    """
    from src.agent.llm import get_llm_safe
    from src.agent.prompts import GENERATE_HTML_PROMPT

    llm = get_llm_safe(temperature=0.3)
    if not llm:
        return _fallback_html(content)

    try:
        prompt = GENERATE_HTML_PROMPT.format(
            document_plan=f"布局: {layout}, 样式: {style}",
            extracted_contents=content[:8000],  # 限制长度防止超上下文
            user_prompt="",
        )
        response = llm.invoke(prompt)
        # 提取 HTML 部分
        html = response.content
        if "```html" in html:
            html = html.split("```html")[1].split("```")[0]
        elif "```" in html:
            html = html.split("```")[1].split("```")[0]
        return html.strip()
    except Exception as e:
        return _fallback_html(content)


@tool
def render_to_pdf(html_content: str, output_filename: str = "generated_doc.pdf") -> str:
    """将 HTML 渲染为 PDF 文件。

    Args:
        html_content: 完整的 HTML 文档字符串
        output_filename: 输出文件名

    Returns:
        生成的 PDF 文件路径
    """
    OUTPUT_DIR.mkdir(exist_ok=True)
    output_path = OUTPUT_DIR / output_filename

    # 方案1: Playwright Chromium 渲染（效果最好，不依赖系统库）
    try:
        from playwright.sync_api import sync_playwright
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            page = browser.new_page()
            page.set_content(html_content, wait_until="networkidle")
            page.pdf(path=str(output_path), format="A4", print_background=True)
            browser.close()
        if output_path.exists():
            return str(output_path)
    except Exception as e:
        print(f"[render_to_pdf] Playwright 渲染失败: {e}")

    # 方案2: 回退到 WeasyPrint
    try:
        from weasyprint import HTML
        HTML(string=html_content).write_pdf(str(output_path))
        return str(output_path)
    except Exception:
        pass

    # 方案3: 保存为 HTML 文件
    html_path = output_path.with_suffix(".html")
    html_path.write_text(html_content, encoding="utf-8")
    return f"PDF 渲染失败，已保存为 HTML: {str(html_path)}"


@tool
def render_to_docx(html_content: str, output_filename: str = "generated_doc.docx") -> str:
    """将 HTML 内容转换为 Word (.docx) 文件。

    Args:
        html_content: 完整的 HTML 文档字符串
        output_filename: 输出文件名

    Returns:
        生成的 Word 文件路径
    """
    OUTPUT_DIR.mkdir(exist_ok=True)
    output_path = OUTPUT_DIR / output_filename

    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # 简单的 HTML → Word 转换
        import re
        # 提取标题
        titles = re.findall(r'<h[1-6][^>]*>(.*?)</h[1-6]>', html_content, re.DOTALL)
        paragraphs = re.findall(r'<p[^>]*>(.*?)</p>', html_content, re.DOTALL)
        # 清理 HTML 标签
        clean = lambda t: re.sub(r'<[^>]+>', '', t).strip()

        if titles:
            doc.add_heading(clean(titles[0]), level=0)

        # 添加所有段落
        for p_text in paragraphs:
            text = clean(p_text)
            if text:
                doc.add_paragraph(text)

        # 如果段落太少，尝试添加所有文本
        if len(paragraphs) < 3:
            all_text = clean(html_content)
            # 按换行分割
            for line in all_text.split('\n'):
                line = line.strip()
                if line and len(line) > 2:
                    doc.add_paragraph(line)

        doc.save(str(output_path))
        return str(output_path)
    except Exception as e:
        return f"Word 生成失败: {str(e)}"


@tool
def search_furniture_glossary(term: str, target_lang: str = "DE") -> str:
    """查询家具行业术语的翻译对照表。

    Args:
        term: 要查询的英文术语
        target_lang: 目标语言代码 (DE/IT/FR/ES/JP)

    Returns:
        术语翻译
    """
    from src.translator.translator import FURNITURE_GLOSSARY
    from src.models import LanguageCode

    term_lower = term.lower()
    results = []
    for en_term, translations in FURNITURE_GLOSSARY.items():
        if term_lower in en_term or en_term in term_lower:
            lang_code = target_lang.upper()
            try:
                lang_enum = LanguageCode(lang_code)
                trans = translations.get(lang_enum, "未找到")
                results.append(f"{en_term} → {trans}")
            except ValueError:
                results.append(f"{en_term} → 不支持的语言: {target_lang}")

    if results:
        return "\n".join(results)
    return f"未找到术语 '{term}' 的翻译"


# === 辅助函数 ===

def _extract_pdf(path: Path) -> str:
    """从 PDF 提取文本"""
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        texts = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                texts.append(f"[第 {i+1} 页]\n{text}")
        return "\n\n".join(texts) if texts else "PDF 无法提取文本（可能是扫描版）"
    except Exception as e:
        return f"PDF 提取失败: {str(e)}"


def _extract_docx(path: Path) -> str:
    """从 Word 文档提取文本"""
    try:
        from docx import Document
        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)
    except Exception as e:
        return f"Word 提取失败: {str(e)}"


def _fallback_html(content: str) -> str:
    """LLM 不可用时的回退 HTML 生成"""
    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <style>
        body {{ font-family: 'Microsoft YaHei', sans-serif; max-width: 800px; margin: 40px auto; padding: 20px; line-height: 1.8; color: #333; }}
        h1 {{ color: #1a1a2e; border-bottom: 2px solid #4361ee; padding-bottom: 10px; }}
        .content {{ white-space: pre-wrap; }}
        .note {{ background: #fff3cd; border: 1px solid #ffc107; padding: 12px; border-radius: 6px; margin-top: 20px; font-size: 13px; }}
    </style>
</head>
<body>
    <h1>Generated Document</h1>
    <div class="content">{content}</div>
    <div class="note">⚠️ 此文档由回退模式生成（LLM 不可用）。请配置 LLM_API_KEY 以启用 AI 排版。</div>
</body>
</html>"""
